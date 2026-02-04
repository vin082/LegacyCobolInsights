"""
Document Generator Agent - Creates detailed technical documentation from knowledge graph
Focuses on program-level functional specifications for developers and business analysts
"""
from typing import Dict, Any, List
from datetime import datetime
from pathlib import Path
from utils.logger import logger
from utils.neo4j_client import neo4j_client
from utils.llm_factory import get_llm
from config.settings import settings
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocumentGeneratorAgent:
    """Agent responsible for generating detailed technical documentation for COBOL programs"""

    def __init__(self):
        self.neo4j = neo4j_client
        self.exports_dir = Path(__file__).parent.parent / "exports"
        self.templates_dir = Path(__file__).parent.parent / "templates"
        self.exports_dir.mkdir(exist_ok=True)
        self.llm = None
        self._embeddings = None
        self._vector_index_available = None  # tri-state: None = unchecked

    # ── Embedding / vector-search plumbing (mirrors cypher_gen pattern) ──

    @property
    def embeddings(self):
        """Lazy-init OpenAI embeddings. Returns None if unavailable."""
        if self._embeddings is None:
            try:
                from langchain_openai import OpenAIEmbeddings
                self._embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
                logger.debug("DocumentGenerator: initialised embeddings")
            except Exception as e:
                logger.warning(f"DocumentGenerator: embeddings init failed: {e}")
                self._embeddings = None
        return self._embeddings

    def _check_vector_index_available(self) -> bool:
        """Check once whether the cobol_program_embeddings index exists."""
        if self._vector_index_available is not None:
            return self._vector_index_available
        try:
            result = self.neo4j.query("SHOW INDEXES WHERE name = 'cobol_program_embeddings'")
            self._vector_index_available = len(result) > 0
            logger.debug(f"DocumentGenerator: vector index available = {self._vector_index_available}")
            return self._vector_index_available
        except Exception:
            self._vector_index_available = False
            return False

    def _semantic_search(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        """
        Vector similarity search against cobol_program_embeddings.
        Returns list of {name, domain, summary, business_logic, score}.
        Silently returns [] if embeddings or index are unavailable.
        """
        if not self.embeddings or not self._check_vector_index_available():
            return []
        try:
            query_embedding = self.embeddings.embed_query(query)
            return self.neo4j.query("""
                CALL db.index.vector.queryNodes('cobol_program_embeddings', $k, $embedding)
                YIELD node, score
                WHERE score > 0.65
                RETURN node.name            AS name,
                       COALESCE(node.domain, 'Unknown') AS domain,
                       node.summary         AS summary,
                       node.business_logic  AS business_logic,
                       score
                ORDER BY score DESC
            """, {'k': k, 'embedding': query_embedding})
        except Exception as e:
            logger.warning(f"DocumentGenerator: semantic search failed: {e}")
            return []

    def _semantic_enrich(self, themes: List[str], k: int = 5) -> str:
        """
        Run a vector search for each theme string and return a single
        formatted context block ready to paste into an LLM prompt.

        If the vector index is unavailable the returned string is empty —
        callers should fall back to the non-vector context they already have.
        """
        if not self.embeddings or not self._check_vector_index_available():
            return ""

        blocks: List[str] = []
        for theme in themes:
            hits = self._semantic_search(theme, k=k)
            if not hits:
                continue
            lines = [f"  - {h['name']} [{h['domain']}] (similarity {h['score']:.2f}): "
                     f"{h.get('summary') or h.get('business_logic') or 'no summary'}"
                     for h in hits]
            blocks.append(f"Theme \"{theme}\":\n" + "\n".join(lines))

        if not blocks:
            return ""
        return "SEMANTICALLY RELEVANT PROGRAMS (vector search):\n" + "\n\n".join(blocks)

    def _get_related_programs(self, program_name: str, domain: str = None,
                              summary: str = None, business_logic: str = None,
                              k: int = 5) -> str:
        """
        Find programs semantically related to *program_name*.

        Strategy:
          1. If the vector index is available AND we have a rich text
             (summary / business_logic), do a vector similarity search
             using that text as the query.  Exclude the program itself.
          2. If vector search returns nothing (program not enriched, or
             index missing), fall back to a lightweight Cypher query that
             finds programs sharing the same domain or copybooks.

        Returns a formatted context string ready to drop into a prompt,
        or empty string if nothing was found.
        """
        vector_hits: List[Dict] = []

        # ── attempt 1: vector search ──
        query_text = (summary or "") + " " + (business_logic or "")
        if query_text.strip() and self.embeddings and self._check_vector_index_available():
            try:
                query_embedding = self.embeddings.embed_query(query_text.strip())
                vector_hits = self.neo4j.query("""
                    CALL db.index.vector.queryNodes('cobol_program_embeddings', $k, $embedding)
                    YIELD node, score
                    WHERE score > 0.6 AND node.name <> $exclude
                    RETURN node.name            AS name,
                           COALESCE(node.domain, 'Unknown') AS domain,
                           node.summary         AS summary,
                           node.business_logic  AS business_logic,
                           score
                    ORDER BY score DESC
                """, {'k': k + 1, 'embedding': query_embedding, 'exclude': program_name})
            except Exception as e:
                logger.warning(f"Vector search for related programs failed: {e}")

        if vector_hits:
            lines = [f"  - {h['name']} [{h['domain']}] (similarity {h['score']:.2f}): "
                     f"{h.get('summary') or h.get('business_logic') or 'no summary'}"
                     for h in vector_hits[:k]]
            return "RELATED PROGRAMS (vector similarity):\n" + "\n".join(lines)

        # ── attempt 2: graph fallback — same domain or shared copybooks ──
        try:
            fallback = self.neo4j.query("""
                MATCH (target:CobolProgram {name: $name})
                OPTIONAL MATCH (target)-[:INCLUDES]->(cb:Copybook)<-[:INCLUDES]-(sibling:CobolProgram)
                  WHERE sibling.name <> $name
                WITH target, collect(DISTINCT sibling.name) AS copybook_siblings
                OPTIONAL MATCH (domain_sibling:CobolProgram)
                  WHERE domain_sibling.domain = target.domain
                    AND domain_sibling.name <> $name
                WITH target, copybook_siblings,
                     collect(DISTINCT domain_sibling.name) AS domain_siblings
                RETURN copybook_siblings, domain_siblings,
                       target.domain AS domain
            """, {'name': program_name})

            if fallback:
                row = fallback[0]
                cb_siblings  = row.get('copybook_siblings') or []
                dom_siblings = row.get('domain_siblings') or []
                lines: List[str] = []
                if cb_siblings:
                    lines.append(f"  Shared copybooks: {', '.join(cb_siblings[:6])}")
                if dom_siblings:
                    lines.append(f"  Same domain ({row.get('domain', 'Unknown')}): {', '.join(dom_siblings[:6])}")
                if lines:
                    return "RELATED PROGRAMS (graph — shared copybooks / domain):\n" + "\n".join(lines)
        except Exception as e:
            logger.warning(f"Fallback related-programs query failed: {e}")

        return ""

    def process(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate comprehensive technical documentation from KG data

        Args:
            state: Dict with doc_type, format, filters, program_name (optional)

        Returns:
            Updated state with file_path and status
        """
        doc_type = state.get('doc_type', 'system_overview')
        doc_format = state.get('format', 'markdown')
        filters = state.get('filters', {})
        program_name = state.get('program_name')  # For single program docs

        logger.info(f"📄 DOCUMENT GENERATOR: Creating {doc_type} in {doc_format}")

        try:
            # Initialize LLM
            self.llm = get_llm(temperature=0.3)

            # Step 1: Gather comprehensive data from Neo4j
            data = self._gather_comprehensive_data(doc_type, filters, program_name)

            # Step 2: Use LLM to generate professional documentation
            content = self._generate_technical_documentation(doc_type, data)

            # Step 3: Export to file (markdown or docx)
            if doc_format == 'markdown':
                file_path = self._export_file(doc_type, 'markdown', content, program_name)
            elif doc_format == 'docx':
                file_path = self._export_as_docx(doc_type, content, program_name)
            else:
                raise ValueError(f"Unsupported format: {doc_format}")

            logger.info(f"✅ Document generated: {file_path}")

            return {
                **state,
                "status": "completed",
                "file_path": str(file_path),
                "stage": "document_generation"
            }

        except Exception as e:
            logger.error(f"Document generation failed: {e}")
            return {
                **state,
                "status": "failed",
                "errors": state.get('errors', []) + [str(e)],
                "stage": "document_generation"
            }

    def _gather_comprehensive_data(self, doc_type: str, filters: Dict, program_name: str = None) -> Dict[str, Any]:
        """
        Query Neo4j for comprehensive program-level data
        """
        logger.info(f"📊 Gathering comprehensive data for {doc_type}")

        data = {}

        if doc_type == 'system_overview':
            # Get all programs for documentation
            data['programs'] = self._get_all_programs_for_documentation(filters)

            # Get system-level stats
            data['stats'] = self.neo4j.get_statistics()
            data['total_programs'] = len(data['programs'])

            # --- Consolidated system-level data ---
            data['call_graph']            = self._get_call_graph_summary()
            data['shared_files']          = self._get_shared_data_files()
            data['shared_copybooks']      = self._get_shared_copybooks()
            data['domain_groups']         = self._get_domain_groups()
            data['call_chains']           = self._get_call_chains()
            data['data_flow_map']         = self._get_data_flow_map()
            data['cross_domain_links']    = self._get_cross_domain_interactions()
            data['program_descriptions']  = self._get_program_descriptions()

        elif doc_type == 'program_detail' and program_name:
            # Get detailed info for single program
            data['program_info'] = self._get_program_details(program_name)
            data['dependencies'] = self._get_program_dependencies(program_name)
            data['data_flows'] = self._get_program_data_flows(program_name)
            data['procedures'] = self._get_program_procedures(program_name)
            data['business_logic'] = self._extract_program_business_logic(program_name)
            data['copybooks'] = self._get_program_copybooks(program_name)
            data['jobs'] = self._get_program_jobs(program_name)

        data['generation_date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        data['filters'] = filters

        return data

    def _get_all_programs_for_documentation(self, filters: Dict) -> List[Dict]:
        """Get all programs with comprehensive details for documentation"""

        # Build WHERE clause based on filters; domain uses a parameter to avoid injection
        where_clauses = []
        params: Dict[str, Any] = {}

        if filters.get('domain'):
            where_clauses.append("p.domain = $domain")
            params['domain'] = filters['domain']

        if filters.get('complexity'):
            if filters['complexity'] == 'low':
                where_clauses.append("p.complexity_score < 30")
            elif filters['complexity'] == 'medium':
                where_clauses.append("p.complexity_score >= 30 AND p.complexity_score < 70")
            elif filters['complexity'] == 'high':
                where_clauses.append("p.complexity_score >= 70")

        where_clause = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""

        # Get limit from filters (default 10 for performance); coerce to int
        limit = int(filters.get('max_programs', 10))
        params['limit'] = limit

        query = f"""
        MATCH (p:CobolProgram)
        {where_clause}
        OPTIONAL MATCH (p)-[:CALLS]->(called:CobolProgram)
        OPTIONAL MATCH (caller:CobolProgram)-[:CALLS]->(p)
        OPTIONAL MATCH (p)-[:CONTAINS]->(proc:Procedure)
        OPTIONAL MATCH (p)-[r]->(f:DataFile)
        WHERE type(r) IN ['READS', 'WRITES']

        WITH p,
             COUNT(DISTINCT called) AS calls_out,
             COUNT(DISTINCT caller) AS calls_in,
             COUNT(DISTINCT proc) AS procedure_count,
             COUNT(DISTINCT CASE WHEN type(r) = 'READS' THEN f END) AS files_read,
             COUNT(DISTINCT CASE WHEN type(r) = 'WRITES' THEN f END) AS files_written,
             CASE WHEN p.code IS NOT NULL THEN size(split(p.code, '\\n')) ELSE 0 END AS estimated_loc

        RETURN p.name AS program_name,
               COALESCE(p.domain, 'Not Enriched') AS domain,
               p.description AS description,
               COALESCE(p.complexity_score, 0) AS complexity,
               COALESCE(p.loc, estimated_loc) AS loc,
               p.code AS source_code,
               calls_out,
               calls_in,
               procedure_count,
               files_read,
               files_written
        ORDER BY COALESCE(p.complexity_score, calls_in + calls_out) DESC
        LIMIT $limit
        """

        try:
            results = self.neo4j.query(query, params)
            logger.info(f"Retrieved {len(results)} programs for documentation")
            return results
        except Exception as e:
            logger.error(f"Error fetching programs: {e}")
            return []

    def _get_program_details(self, program_name: str) -> Dict[str, Any]:
        """Get detailed information for a specific program, including source code"""
        query = """
        MATCH (p:CobolProgram {name: $program_name})
        RETURN p.name AS name,
               p.domain AS domain,
               p.description AS description,
               p.complexity_score AS complexity_score,
               p.loc AS loc,
               p.file_path AS file_path,
               p.code AS source_code
        """
        try:
            result = self.neo4j.query(query, {'program_name': program_name})
            return result[0] if result else {}
        except Exception as e:
            logger.error(f"Error fetching details for {program_name}: {e}")
            return {}

    def _get_program_dependencies(self, program_name: str) -> Dict[str, Any]:
        """Get program dependencies (calls and called by)"""
        params = {'program_name': program_name}

        # Programs this program calls
        calls_query = """
        MATCH (p:CobolProgram {name: $program_name})-[:CALLS]->(called:CobolProgram)
        RETURN called.name AS program,
               called.domain AS domain,
               called.complexity_score AS complexity
        ORDER BY called.name
        """

        # Programs that call this program
        called_by_query = """
        MATCH (caller:CobolProgram)-[:CALLS]->(p:CobolProgram {name: $program_name})
        RETURN caller.name AS program,
               caller.domain AS domain,
               caller.complexity_score AS complexity
        ORDER BY caller.name
        """

        try:
            calls = self.neo4j.query(calls_query, params)
            called_by = self.neo4j.query(called_by_query, params)
            return {
                'calls': calls,
                'called_by': called_by
            }
        except Exception as e:
            logger.error(f"Error fetching dependencies for {program_name}: {e}")
            return {'calls': [], 'called_by': []}

    def _get_program_data_flows(self, program_name: str) -> Dict[str, Any]:
        """Get data file operations for a program"""
        query = """
        MATCH (p:CobolProgram {name: $program_name})-[r]->(f:DataFile)
        WHERE type(r) IN ['READS', 'WRITES']
        RETURN f.name AS file_name,
               type(r) AS operation,
               f.description AS file_description
        ORDER BY f.name, operation
        """

        try:
            results = self.neo4j.query(query, {'program_name': program_name})

            # Organize by file
            files = {}
            for row in results:
                file_name = row['file_name']
                if file_name not in files:
                    files[file_name] = {
                        'description': row.get('file_description', ''),
                        'operations': []
                    }
                files[file_name]['operations'].append(row['operation'])

            return files
        except Exception as e:
            logger.error(f"Error fetching data flows for {program_name}: {e}")
            return {}

    def _get_program_procedures(self, program_name: str) -> List[Dict]:
        """Get procedures/paragraphs within a program"""
        query = """
        MATCH (p:CobolProgram {name: $program_name})-[:CONTAINS]->(proc:Procedure)
        RETURN proc.name AS name,
               proc.type AS type,
               proc.description AS description
        ORDER BY proc.name
        LIMIT 100
        """

        try:
            return self.neo4j.query(query, {'program_name': program_name})
        except Exception as e:
            logger.error(f"Error fetching procedures for {program_name}: {e}")
            return []

    def _extract_program_business_logic(self, program_name: str) -> str:
        """Extract business logic description from enrichment data"""
        query = """
        MATCH (p:CobolProgram {name: $program_name})
        RETURN p.business_logic AS business_logic,
               p.description AS description
        """

        try:
            result = self.neo4j.query(query, {'program_name': program_name})
            if result:
                return result[0].get('business_logic') or result[0].get('description') or ''
            return ''
        except Exception as e:
            logger.error(f"Error fetching business logic for {program_name}: {e}")
            return ''

    def _get_program_copybooks(self, program_name: str) -> List[Dict]:
        """Get copybooks included by a program"""
        query = """
        MATCH (p:CobolProgram {name: $program_name})-[:INCLUDES]->(c:Copybook)
        RETURN c.name AS name
        ORDER BY c.name
        """
        try:
            return self.neo4j.query(query, {'program_name': program_name})
        except Exception as e:
            logger.error(f"Error fetching copybooks for {program_name}: {e}")
            return []

    def _get_program_jobs(self, program_name: str) -> List[Dict]:
        """Get JCL jobs that execute a program"""
        query = """
        MATCH (j:Job)-[:EXECUTES]->(p:CobolProgram {name: $program_name})
        RETURN j.name AS name,
               j.description AS description
        ORDER BY j.name
        """
        try:
            return self.neo4j.query(query, {'program_name': program_name})
        except Exception as e:
            logger.error(f"Error fetching jobs for {program_name}: {e}")
            return []

    # ── System-level consolidated queries ────────────────────────────────

    def _get_call_graph_summary(self) -> Dict[str, Any]:
        """
        Return hub/leaf/entry-point info for the call graph.
        - hub_callers:  programs that CALL the most other programs
        - hub_callees:  programs that are CALLED BY the most other programs
        - entry_points: programs that nobody calls (top-level / batch drivers)
        """
        hub_callers_q = """
        MATCH (p:CobolProgram)-[:CALLS]->(c:CobolProgram)
        WITH p, count(c) AS out_count
        WHERE out_count > 0
        RETURN p.name AS program, out_count AS calls_out,
               COALESCE(p.domain, 'Unknown') AS domain
        ORDER BY out_count DESC
        LIMIT 15
        """
        hub_callees_q = """
        MATCH (caller:CobolProgram)-[:CALLS]->(p:CobolProgram)
        WITH p, count(caller) AS in_count
        WHERE in_count > 0
        RETURN p.name AS program, in_count AS calls_in,
               COALESCE(p.domain, 'Unknown') AS domain
        ORDER BY in_count DESC
        LIMIT 15
        """
        entry_points_q = """
        MATCH (p:CobolProgram)
        WHERE NOT EXISTS { MATCH (:CobolProgram)-[:CALLS]->(p) }
          AND EXISTS { MATCH (p)-[:CALLS]->(:CobolProgram) }
        RETURN p.name AS program,
               COALESCE(p.domain, 'Unknown') AS domain,
               COALESCE(p.complexity_score, 0) AS complexity
        ORDER BY p.name
        LIMIT 20
        """
        try:
            return {
                'hub_callers':  self.neo4j.query(hub_callers_q),
                'hub_callees':  self.neo4j.query(hub_callees_q),
                'entry_points': self.neo4j.query(entry_points_q),
            }
        except Exception as e:
            logger.error(f"Error fetching call graph summary: {e}")
            return {'hub_callers': [], 'hub_callees': [], 'entry_points': []}

    def _get_shared_data_files(self) -> List[Dict]:
        """
        Data files accessed by more than one program — the shared / critical
        data layer of the system.
        """
        query = """
        MATCH (p:CobolProgram)-[r]->(f:DataFile)
        WHERE type(r) IN ['READS', 'WRITES']
        WITH f, collect(DISTINCT p.name) AS programs, collect(DISTINCT type(r)) AS ops
        WHERE size(programs) > 1
        RETURN f.name AS file_name,
               programs AS used_by,
               ops AS operations,
               size(programs) AS program_count
        ORDER BY program_count DESC
        LIMIT 20
        """
        try:
            return self.neo4j.query(query)
        except Exception as e:
            logger.error(f"Error fetching shared data files: {e}")
            return []

    def _get_shared_copybooks(self) -> List[Dict]:
        """Copybooks included by more than one program — shared data structures."""
        query = """
        MATCH (p:CobolProgram)-[:INCLUDES]->(c:Copybook)
        WITH c, collect(DISTINCT p.name) AS programs
        WHERE size(programs) > 1
        RETURN c.name AS copybook,
               programs AS used_by,
               size(programs) AS program_count
        ORDER BY program_count DESC
        LIMIT 20
        """
        try:
            return self.neo4j.query(query)
        except Exception as e:
            logger.error(f"Error fetching shared copybooks: {e}")
            return []

    def _get_domain_groups(self) -> List[Dict]:
        """Programs grouped by domain with aggregate metrics."""
        query = """
        MATCH (p:CobolProgram)
        WITH COALESCE(p.domain, 'Unknown') AS domain,
             collect(p.name) AS programs,
             count(*) AS count,
             round(avg(COALESCE(p.complexity_score, 0)), 1) AS avg_complexity,
             sum(COALESCE(p.loc, 0)) AS total_loc
        RETURN domain, programs, count, avg_complexity, total_loc
        ORDER BY count DESC
        """
        try:
            return self.neo4j.query(query)
        except Exception as e:
            logger.error(f"Error fetching domain groups: {e}")
            return []

    def _get_call_chains(self) -> List[Dict]:
        """
        Walk CALLS edges up to depth 3 from every entry point.
        Returns rows: caller → called, with domains, so the LLM can
        reconstruct end-to-end business flows.
        """
        query = """
        MATCH (entry:CobolProgram)
        WHERE NOT EXISTS { MATCH (:CobolProgram)-[:CALLS]->(entry) }
        MATCH chain = (entry)-[:CALLS*1..3]->(leaf:CobolProgram)
        WITH entry, leaf, [n IN nodes(chain) | n.name] AS path,
             [n IN nodes(chain) | COALESCE(n.domain, 'Unknown')] AS domains
        RETURN entry.name AS entry_point,
               COALESCE(entry.domain, 'Unknown') AS entry_domain,
               path AS call_path,
               domains AS path_domains,
               leaf.name AS leaf_program
        LIMIT 60
        """
        try:
            return self.neo4j.query(query)
        except Exception as e:
            logger.error(f"Error fetching call chains: {e}")
            return []

    def _get_data_flow_map(self) -> List[Dict]:
        """
        For every data file: which programs READ it and which WRITE it.
        Gives the LLM enough to narrate input → processing → output flows.
        """
        query = """
        MATCH (p:CobolProgram)-[r]->(f:DataFile)
        WHERE type(r) IN ['READS', 'WRITES']
        WITH f,
             [p2 IN [(p2)-[:READS]->(f) | p2.name] | p2] AS readers,
             [p2 IN [(p2)-[:WRITES]->(f) | p2.name] | p2] AS writers
        RETURN f.name AS file_name,
               readers AS readers,
               writers AS writers
        ORDER BY size(readers) + size(writers) DESC
        LIMIT 25
        """
        try:
            return self.neo4j.query(query)
        except Exception as e:
            logger.error(f"Error fetching data flow map: {e}")
            return []

    def _get_cross_domain_interactions(self) -> List[Dict]:
        """
        CALLS relationships where caller and callee are in different domains.
        These are the seams / coupling points between business areas.
        """
        query = """
        MATCH (a:CobolProgram)-[:CALLS]->(b:CobolProgram)
        WHERE COALESCE(a.domain, 'Unknown') <> COALESCE(b.domain, 'Unknown')
        RETURN a.name AS caller,
               COALESCE(a.domain, 'Unknown') AS caller_domain,
               b.name AS callee,
               COALESCE(b.domain, 'Unknown') AS callee_domain
        ORDER BY a.name
        LIMIT 30
        """
        try:
            return self.neo4j.query(query)
        except Exception as e:
            logger.error(f"Error fetching cross-domain interactions: {e}")
            return []

    def _get_program_descriptions(self) -> List[Dict]:
        """
        Collect business_logic + description fields from all programs.
        Used by the LLM to extract and consolidate key business rules.
        """
        query = """
        MATCH (p:CobolProgram)
        WHERE p.business_logic IS NOT NULL OR p.description IS NOT NULL
        RETURN p.name AS name,
               COALESCE(p.domain, 'Unknown') AS domain,
               p.business_logic AS business_logic,
               p.description AS description
        ORDER BY p.name
        LIMIT 50
        """
        try:
            return self.neo4j.query(query)
        except Exception as e:
            logger.error(f"Error fetching program descriptions: {e}")
            return []

    def _generate_technical_documentation(self, doc_type: str, data: Dict) -> str:
        """Generate technical documentation using LLM"""

        if doc_type == 'system_overview':
            return self._generate_system_documentation(data)
        elif doc_type == 'program_detail':
            return self._generate_program_documentation(data)
        else:
            raise ValueError(f"Unknown doc_type: {doc_type}")

    # ── LLM helper: invoke with a fallback message ──────────────────────
    def _llm_or_fallback(self, prompt: str, fallback: str) -> str:
        try:
            return self.llm.invoke(prompt).content.strip()
        except Exception as e:
            logger.error(f"LLM call failed: {e}")
            return fallback

    def _generate_system_documentation(self, data: Dict) -> str:
        """
        Generate a consolidated system-level document.
        Structure:
          1.  Architecture Narrative          (LLM)
          2.  End-to-End Business Flows       (LLM — new)
          3.  Data Flow Summary               (LLM — new)
          4.  Risk & Impact                   (LLM — new)
          5.  Key Business Rules              (LLM — new)
          6.  System Statistics
          7.  Domain Breakdown
          8.  Call-Graph Analysis
          9.  Shared Data Layer
          10. Shared Copybooks
          11. Complexity Hotspots
          12. Program Catalog
        """

        programs             = data.get('programs', [])
        stats                = data.get('stats', {})
        call_graph           = data.get('call_graph', {})
        shared_files         = data.get('shared_files', [])
        shared_cbs           = data.get('shared_copybooks', [])
        domain_groups        = data.get('domain_groups', [])
        call_chains          = data.get('call_chains', [])
        data_flow_map        = data.get('data_flow_map', [])
        cross_domain_links   = data.get('cross_domain_links', [])
        program_descriptions = data.get('program_descriptions', [])

        enriched_count      = sum(1 for p in programs if p.get('domain') not in ['Not Enriched', None])
        not_enriched_count  = len(programs) - enriched_count
        total_loc           = sum(p.get('loc') or 0 for p in programs)
        total_relationships = stats.get('total_relationships', 0)

        entry_pts    = call_graph.get('entry_points', [])
        hub_callers  = call_graph.get('hub_callers', [])
        hub_callees  = call_graph.get('hub_callees', [])

        # ═══════════════════════════════════════════════════════════════
        # Pre-format context strings reused across multiple prompts
        # ═══════════════════════════════════════════════════════════════
        domain_summary = "\n".join(
            f"  - {g['domain']}: {g['count']} programs, avg complexity {g['avg_complexity']}, {g['total_loc']} total LOC"
            for g in domain_groups
        ) or "  - No domain data available"

        call_graph_summary = (
            f"Entry points (called by nobody): {[e['program'] for e in entry_pts[:10]]}\n"
            f"Most outbound calls: {[(h['program'], h['calls_out']) for h in hub_callers[:5]]}\n"
            f"Most inbound calls:  {[(h['program'], h['calls_in'])  for h in hub_callees[:5]]}"
        )

        shared_files_summary = "\n".join(
            f"  - {f['file_name']}: used by {f['program_count']} programs ({', '.join(f['used_by'][:5])})"
            for f in shared_files[:10]
        ) or "  - No shared data files found"

        shared_cbs_summary = "\n".join(
            f"  - {c['copybook']}: included by {c['program_count']} programs"
            for c in shared_cbs[:10]
        ) or "  - No shared copybooks found"

        call_chains_text = "\n".join(
            f"  {' → '.join(row['call_path'])}   (domains: {' → '.join(row['path_domains'])})"
            for row in call_chains[:30]
        ) or "  - No call chains found"

        data_flow_text = "\n".join(
            f"  - {row['file_name']}:  written by {row['writers'] or ['(none)']}  |  read by {row['readers'] or ['(none)']}"
            for row in data_flow_map[:20]
        ) or "  - No data flow information found"

        cross_domain_text = "\n".join(
            f"  - {row['caller']} [{row['caller_domain']}] → {row['callee']} [{row['callee_domain']}]"
            for row in cross_domain_links[:20]
        ) or "  - No cross-domain calls detected"

        descriptions_text = "\n".join(
            f"  - {row['name']} [{row['domain']}]: {row.get('business_logic') or row.get('description') or 'N/A'}"
            for row in program_descriptions[:40]
        ) or "  - No program descriptions available"

        # ═══════════════════════════════════════════════════════════════
        # LLM calls — 5 focused prompts, each with a single job
        # ═══════════════════════════════════════════════════════════════

        # 1. Architecture Narrative
        narrative = self._llm_or_fallback(
            f"""You are a technical architect documenting a legacy COBOL system.
Write a concise but insightful **architecture narrative** (4-6 paragraphs) that explains:
- What this system does as a whole (infer from domain groups and program names)
- How programs are organised by domain/function
- How the call graph connects them (entry points → hubs → leaves)
- What the shared data files and copybooks tell us about coupling
- Any structural risks or observations (e.g. tightly-coupled hubs, orphan programs)

Use ONLY the facts below. Do not invent program names or relationships.

DOMAIN GROUPS:
{domain_summary}

CALL GRAPH:
{call_graph_summary}

SHARED DATA FILES (accessed by >1 program):
{shared_files_summary}

SHARED COPYBOOKS (included by >1 program):
{shared_cbs_summary}

TOTAL PROGRAMS: {len(programs)} | TOTAL LOC: {total_loc} | TOTAL RELATIONSHIPS: {total_relationships}
""",
            "Architecture narrative could not be generated. See the sections below for system-level details."
        )

        # ── Vector-search enrichment for sections 2 & 5 ──────────────────
        # One call per theme set; results are spliced into the prompts only
        # when the vector index is available — otherwise the prompts work
        # fine with the graph-based context alone.
        flow_vector_ctx = self._semantic_enrich([
            "transaction processing flow",
            "account enquiry and lookup",
            "batch settlement and reconciliation",
            "card validation and verification",
            "customer and account management",
        ], k=4)

        rules_vector_ctx = self._semantic_enrich([
            "validation and error handling rules",
            "account and balance rules",
            "transaction posting and limits",
            "card and customer verification",
            "file processing and data integrity",
        ], k=4)

        # 2. End-to-End Business Flows
        business_flows = self._llm_or_fallback(
            f"""You are a business analyst documenting a legacy COBOL system for stakeholders.
Using the call chains and program summaries below, identify and describe 3-5 end-to-end business flows (e.g. "Transaction Processing", "Account Enquiry", "Batch Settlement").

For each flow:
- Give it a clear business name
- Describe what happens in plain language (no code)
- List the programs involved in order
- Note which domain each step belongs to

Use ONLY the data below. Infer business meaning from program names and domain labels. Do not invent programs.

CALL CHAINS (entry point → ... → leaf, with domains):
{call_chains_text}

DOMAIN GROUPS:
{domain_summary}
{(chr(10) + flow_vector_ctx) if flow_vector_ctx else ''}
""",
            "Business flow analysis could not be generated. Refer to the Call-Graph Analysis section for program relationships."
        )

        # 3. Data Flow Summary
        data_flow_summary = self._llm_or_fallback(
            f"""You are a business analyst documenting data movement in a legacy COBOL system.
Using the data below, describe how data flows through the system in plain language.

For each significant data file:
- What does it represent (infer from the file name)?
- Which programs produce/write data into it?
- Which programs consume/read data from it?
- What does that tell us about the processing sequence (input → processing → output)?

Group related files into logical flows where possible. Write 3-5 paragraphs. Do not invent file names or programs.

DATA FLOW MAP (file: writers → readers):
{data_flow_text}

SHARED DATA FILES (accessed by >1 program):
{shared_files_summary}
""",
            "Data flow analysis could not be generated. Refer to the Shared Data Layer section for file-level details."
        )

        # 4. Risk & Impact
        risk_summary = self._llm_or_fallback(
            f"""You are a risk analyst reviewing a legacy COBOL system for a stakeholder briefing.
Identify and explain the top risks based on the facts below. For each risk:
- Name the risk clearly
- Explain why it matters to the business
- Identify the specific programs or files involved
- Suggest what should be investigated or mitigated

Focus on: single points of failure, tightly-coupled domains, high-complexity programs that many others depend on, and shared data files that if corrupted would affect many processes.

Write as a prioritised list (highest risk first). Be concise — 1-3 sentences per risk. Do not invent data.

HUB PROGRAMS (most called — if these fail, many processes stop):
{chr(10).join(f"  - {h['program']}: called by {h['calls_in']} programs [{h['domain']}]" for h in hub_callees[:10]) or "  - None"}

CROSS-DOMAIN CALLS (coupling between business areas):
{cross_domain_text}

SHARED DATA FILES:
{shared_files_summary}

HIGH-COMPLEXITY PROGRAMS (score ≥ 70):
{chr(10).join(f"  - {p.get('program_name')}: complexity {p.get('complexity') or 0}, {p.get('loc') or 0} LOC" for p in sorted(programs, key=lambda x: x.get('complexity') or 0, reverse=True)[:10] if (p.get('complexity') or 0) >= 70) or "  - None"}
""",
            "Risk analysis could not be generated. Refer to the Complexity Hotspots and Shared Data Layer sections."
        )

        # 5. Key Business Rules
        business_rules = self._llm_or_fallback(
            f"""You are a business analyst extracting key business rules from a legacy COBOL system.
Read the program descriptions and business logic below. Extract and consolidate the distinct business rules the system enforces.

Group rules by theme (e.g. "Validation Rules", "Account Rules", "Transaction Rules").
For each rule:
- State the rule clearly in plain business language
- Note which program(s) enforce it

Write as a grouped list. Use ONLY information from the descriptions below. Do not invent rules.

PROGRAM DESCRIPTIONS AND BUSINESS LOGIC:
{descriptions_text}
{(chr(10) + rules_vector_ctx) if rules_vector_ctx else ''}
""",
            "Business rules extraction could not be generated. Refer to individual program documentation for rule details."
        )

        # ═══════════════════════════════════════════════════════════════
        # Assemble the markdown document
        # ═══════════════════════════════════════════════════════════════

        enrichment_warning = ""
        if not_enriched_count > 0:
            enrichment_warning = (
                f"\n> ⚠️ **ENRICHMENT STATUS**: {not_enriched_count} of {len(programs)} programs have not been "
                f"enriched yet. Run the enrichment agent to populate domain, complexity, and business-logic metadata.\n"
            )

        md = f"""# COBOL System — Consolidated Technical Overview

**Document Type:** System Architecture & Reference Guide
**Generated:** {data['generation_date']}
**Programs in scope:** {len(programs)}  |  **Total LOC:** {total_loc}  |  **Enriched:** {enriched_count}/{len(programs)}
{enrichment_warning}
---

## Table of Contents

1. [Architecture Narrative](#1-architecture-narrative)
2. [End-to-End Business Flows](#2-end-to-end-business-flows)
3. [Data Flow Summary](#3-data-flow-summary)
4. [Risk & Impact](#4-risk--impact)
5. [Key Business Rules](#5-key-business-rules)
6. [System Statistics](#6-system-statistics)
7. [Domain Breakdown](#7-domain-breakdown)
8. [Call-Graph Analysis](#8-call-graph-analysis)
9. [Shared Data Layer](#9-shared-data-layer)
10. [Shared Copybooks](#10-shared-copybooks)
11. [Complexity Hotspots](#11-complexity-hotspots)
12. [Program Catalog](#12-program-catalog)

---

## 1. Architecture Narrative

{narrative}

---

## 2. End-to-End Business Flows

{business_flows}

---

## 3. Data Flow Summary

{data_flow_summary}

---

## 4. Risk & Impact

{risk_summary}

---

## 5. Key Business Rules

{business_rules}

---

## 6. System Statistics

| Metric | Value |
|--------|-------|
| Total Programs (in graph) | {stats['nodes'].get('CobolProgram', 0)} |
| Programs in this report | {len(programs)} |
| Total Lines of Code | {total_loc} |
| Total Data Files | {stats['nodes'].get('DataFile', 0)} |
| Total Copybooks | {stats['nodes'].get('Copybook', 0)} |
| Total Procedures | {stats['nodes'].get('Procedure', 0)} |
| Total Relationships | {total_relationships} |
| Shared Data Files (>1 program) | {len(shared_files)} |
| Shared Copybooks (>1 program) | {len(shared_cbs)} |
| Entry-point Programs | {len(entry_pts)} |

---

## 7. Domain Breakdown

"""
        if domain_groups:
            md += "| Domain | Programs | Avg Complexity | Total LOC |\n"
            md += "|--------|----------|----------------|-----------|\n"
            for g in domain_groups:
                md += f"| {g['domain']} | {g['count']} | {g['avg_complexity']} | {g['total_loc']} |\n"
            md += "\n"
            for g in domain_groups:
                prog_list = ", ".join(f"`{p}`" for p in sorted(g['programs']))
                md += f"**{g['domain']}:** {prog_list}\n\n"
        else:
            md += "No domain information available.\n\n"

        # ── 8. Call graph ──────────────────────────────────────────────
        md += "---\n\n## 8. Call-Graph Analysis\n\n"

        if entry_pts:
            md += "### Entry Points\n\nThese programs are not called by any other program — they are likely batch drivers or top-level orchestrators.\n\n"
            md += "| Program | Domain | Complexity |\n|---------|--------|------------|\n"
            for e in entry_pts:
                md += f"| `{e['program']}` | {e['domain']} | {int(e['complexity'])} |\n"
            md += "\n"
        else:
            md += "### Entry Points\n\nNo entry points detected (every program is called by at least one other).\n\n"

        if hub_callers:
            md += "### Most Outbound Calls (orchestrators / drivers)\n\n"
            md += "| Program | Calls Out | Domain |\n|---------|-----------|--------|\n"
            for h in hub_callers:
                md += f"| `{h['program']}` | {h['calls_out']} | {h['domain']} |\n"
            md += "\n"

        if hub_callees:
            md += "### Most Inbound Calls (shared services / utilities)\n\n"
            md += "| Program | Called By | Domain |\n|---------|----------|--------|\n"
            for h in hub_callees:
                md += f"| `{h['program']}` | {h['calls_in']} | {h['domain']} |\n"
            md += "\n"

        # ── 9. Shared data files ──────────────────────────────────────
        md += "---\n\n## 9. Shared Data Layer\n\n"
        if shared_files:
            md += "These data files are accessed by more than one program. They represent the shared state of the system and are key integration points.\n\n"
            md += "| Data File | Programs Using It | Operations |\n|-----------|-------------------|-------------|\n"
            for f in shared_files:
                progs = ", ".join(f"`{p}`" for p in f['used_by'])
                ops   = ", ".join(f['operations'])
                md += f"| `{f['file_name']}` | {progs} | {ops} |\n"
            md += "\n"
        else:
            md += "No shared data files detected — programs appear to operate on independent data sets.\n\n"

        # ── 10. Shared copybooks ──────────────────────────────────────
        md += "---\n\n## 10. Shared Copybooks\n\n"
        if shared_cbs:
            md += "Copybooks included by multiple programs define shared data structures across the system.\n\n"
            md += "| Copybook | Included By | Program Count |\n|----------|-------------|---------------|\n"
            for c in shared_cbs:
                progs = ", ".join(f"`{p}`" for p in c['used_by'])
                md += f"| `{c['copybook']}` | {progs} | {c['program_count']} |\n"
            md += "\n"
        else:
            md += "No shared copybooks detected.\n\n"

        # ── 11. Complexity hotspots ───────────────────────────────────
        md += "---\n\n## 11. Complexity Hotspots\n\n"
        high_complexity = [p for p in sorted(programs, key=lambda p: p.get('complexity') or 0, reverse=True)
                           if (p.get('complexity') or 0) >= 70][:10]
        if high_complexity:
            md += "The following programs have a complexity score ≥ 70 and should be prioritised for refactoring or test coverage.\n\n"
            md += "| Program | Domain | Complexity | LOC | Calls Out | Called By |\n"
            md += "|---------|--------|------------|-----|-----------|----------|\n"
            for p in high_complexity:
                md += (f"| `{p.get('program_name', 'N/A')}` | {p.get('domain', 'N/A')} "
                       f"| {p.get('complexity') or 0} | {p.get('loc') or 0} "
                       f"| {p.get('calls_out') or 0} | {p.get('calls_in') or 0} |\n")
            md += "\n"
        else:
            md += "No programs exceed the complexity threshold of 70.\n\n"

        # ── 12. Program catalog (quick-reference only) ───────────────
        md += "---\n\n## 12. Program Catalog\n\n"
        md += "Quick-reference table of all programs in scope. For detailed per-program documentation use the **Single Program** mode.\n\n"
        md += "| # | Program | Domain | Complexity | LOC | Calls Out | Called By | Files R | Files W |\n"
        md += "|---|---------|--------|------------|-----|-----------|----------|---------|----------|\n"
        for i, prog in enumerate(programs, 1):
            md += (f"| {i} | `{prog.get('program_name', 'N/A')}` | {prog.get('domain', 'N/A')} "
                   f"| {prog.get('complexity') or 0} | {prog.get('loc') or 0} "
                   f"| {prog.get('calls_out') or 0} | {prog.get('calls_in') or 0} "
                   f"| {prog.get('files_read') or 0} | {prog.get('files_written') or 0} |\n")

        md += f"""

---

**END OF DOCUMENT**

*Generated by COBOL Agentic Knowledge Graph — {data['generation_date']}*
*For per-program deep dives use Single Program mode.*
"""
        return md

    @staticmethod
    def _extract_paragraphs_from_source(source_code: str) -> List[Dict]:
        """
        Fallback: extract paragraph names from source code when the graph has no
        Procedure nodes.  Strategy: collect all PERFORM targets — these are the
        authoritative set of paragraph/section names actually invoked in the program.
        Filters out COBOL keywords that can follow PERFORM (e.g. UNTIL).
        """
        if not source_code:
            return []

        # Words that can legally follow PERFORM but are not paragraph names
        COBOL_KEYWORDS = {
            'UNTIL', 'VARYING', 'THRU', 'THROUGH', 'INLINE',
            'WITH', 'TEST', 'AFTER', 'BEFORE',
        }

        seen = set()
        paragraphs = []
        # Match PERFORM <name> — name is alphanumeric + hyphens
        for match in re.finditer(r'\bPERFORM\s+([\w-]+)', source_code, re.IGNORECASE):
            name = match.group(1).strip()
            upper = name.upper()
            if upper in COBOL_KEYWORDS or upper in seen:
                continue
            seen.add(upper)
            paragraphs.append({'name': name, 'type': 'Paragraph', 'description': None})

        return paragraphs

    def _generate_single_program_doc(self, prog_summary: Dict, index: int,
                                     copybooks: List[Dict] = None,
                                     jobs: List[Dict] = None) -> str:
        """Generate detailed documentation for a single program using LLM"""

        program_name = prog_summary.get('program_name', 'Unknown')

        # Get detailed data for this program
        prog_details = self._get_program_details(program_name)
        dependencies = self._get_program_dependencies(program_name)
        data_flows = self._get_program_data_flows(program_name)
        procedures = self._get_program_procedures(program_name)
        business_logic = self._extract_program_business_logic(program_name)

        # Source code: prefer prog_details (program_detail path), fall back to prog_summary (system_overview path)
        source_code = prog_details.get('source_code') or prog_summary.get('source_code') or ''
        code_snippet = source_code[:2000] if source_code else ''

        # If graph returned no procedures but we have source, extract paragraph names as fallback
        if not procedures and source_code:
            procedures = self._extract_paragraphs_from_source(source_code)
            logger.info(f"📝 Extracted {len(procedures)} paragraphs from source for {program_name}")

        # Check if enrichment has run
        is_enriched = prog_summary.get('domain') not in ['Not Enriched', None]

        # Build context based on what's available
        if is_enriched:
            context = f"""PROGRAM: {program_name}
DOMAIN: {prog_summary.get('domain', 'Unknown')}
DESCRIPTION: {prog_details.get('description', 'No description available')}
BUSINESS LOGIC: {business_logic or 'Not documented'}
COMPLEXITY: {prog_summary.get('complexity', 0)}
LINES OF CODE: {prog_summary.get('loc', 0)}"""
        else:
            context = f"""PROGRAM: {program_name}
DESCRIPTION: {prog_details.get('description', 'No description available')}
COMPLEXITY: {prog_summary.get('complexity', 0)}
LINES OF CODE: {prog_summary.get('loc', 0)}

NOTE: This program has not been enriched yet. Analyze the source code below to infer:
- Business domain
- Complexity level
- Business logic

SOURCE CODE SNIPPET (first 2000 characters):
```cobol
{code_snippet}
```"""

        # --- Rich data-file context for the LLM prompt ---
        data_files_context = ""
        if data_flows:
            data_files_context = "DATA FILES (from knowledge graph):\n"
            for fname, finfo in list(data_flows.items())[:10]:
                ops = ", ".join(set(finfo['operations']))
                desc = finfo.get('description') or 'No description'
                data_files_context += f"  - {fname}: operations=[{ops}], description={desc}\n"
        # If source is available, also note files referenced in source that may not be in the graph
        if source_code:
            source_upper = source_code.upper()
            select_files = re.findall(r'SELECT\s+([\w-]+)\s+ASSIGN', source_upper)
            graph_files_upper = {f.upper() for f in data_flows.keys()} if data_flows else set()
            extra_files = [f for f in select_files if f not in graph_files_upper]
            if extra_files:
                data_files_context += "\nAdditional files referenced in source (not tracked as graph edges):\n"
                for f in extra_files:
                    data_files_context += f"  - {f}\n"

        # --- Copybooks context ---
        copybooks_context = ""
        if copybooks:
            copybooks_context = f"COPYBOOKS INCLUDED ({len(copybooks)}):\n"
            copybooks_context += "  " + ", ".join(c['name'] for c in copybooks) + "\n"

        # --- JCL jobs context ---
        jobs_context = ""
        if jobs:
            jobs_context = f"JCL JOBS THAT EXECUTE THIS PROGRAM ({len(jobs)}):\n"
            for j in jobs:
                jobs_context += f"  - {j['name']}: {j.get('description') or 'No description'}\n"
        else:
            jobs_context = "JCL JOBS: None found in knowledge graph.\n"

        # --- Procedures context ---
        procedures_context = f"PROCEDURES / PARAGRAPHS ({len(procedures)}):\n"
        if procedures:
            procedures_context += "  " + ", ".join(p['name'] for p in procedures[:30]) + "\n"
            if len(procedures) > 30:
                procedures_context += f"  ...and {len(procedures) - 30} more\n"
        else:
            procedures_context += "  None found.\n"

        # --- Complexity flag for prompt ---
        complexity = prog_summary.get('complexity') or 0
        complexity_instruction = ""
        if complexity >= 70:
            complexity_instruction = f"\nIMPORTANT: This program has a HIGH complexity score of {complexity}. Flag this prominently in the Technical Notes section — explain likely complexity drivers and recommend refactoring or test coverage.\n"

        # --- Related programs via vector search (+ graph fallback) ---
        related_programs_context = self._get_related_programs(
            program_name,
            domain=prog_summary.get('domain'),
            summary=prog_details.get('summary'),
            business_logic=business_logic,
        )

        # Use LLM to generate comprehensive documentation
        prompt = f"""You are a technical writer creating detailed COBOL program documentation for developers and business analysts.

Generate comprehensive documentation for this COBOL program:

{context}

CALLS OUT TO ({len(dependencies.get('calls', []))} programs):
{[p['program'] for p in dependencies.get('calls', [])[:10]]}

CALLED BY ({len(dependencies.get('called_by', []))} programs):
{[p['program'] for p in dependencies.get('called_by', [])[:10]]}

{data_files_context}
{procedures_context}
{copybooks_context}
{jobs_context}
{related_programs_context}
{complexity_instruction}
Generate documentation with these sections:

### 3.{index} {program_name}

#### Purpose and Overview
[2-3 sentences describing what this program does and its role in the system]

#### Business Logic
[Detailed explanation of the business rules and processes this program implements]

#### Key Functionality
[Bullet points of main functions/capabilities]

#### Data Operations
[Describe ALL data files it reads/writes and why — include both graph-tracked and source-referenced files. Use a table with columns: File, Organization/Access, Role]

#### Dependencies
[Explain programs it depends on and which depend on it. Note copybooks. Note any external runtime calls like CEE3ABD if present in source. Note JCL job linkage. If RELATED PROGRAMS context is provided, add a brief "Related Programs" sub-bullet noting programs with similar functionality and why they are relevant.]

#### Technical Notes
[Complexity concerns, performance considerations, maintenance notes. If complexity is high, flag prominently with specific drivers.]

Format as markdown. Be technical but clear. Focus on helping developers understand how to work with this program."""

        try:
            response = self.llm.invoke(prompt)
            llm_output = response.content.strip()

            # Append complexity warning block after LLM output if high complexity
            if complexity >= 70:
                llm_output += f"""

> **HIGH COMPLEXITY WARNING**: This program has a complexity score of {complexity}/100. Prioritize it for refactoring or comprehensive test coverage before any modifications."""

            return llm_output
        except Exception as e:
            logger.error(f"LLM documentation generation failed for {program_name}: {e}")

            # Fallback: Generate basic documentation without LLM
            return self._generate_basic_program_doc(program_name, prog_summary, prog_details,
                                                    dependencies, data_flows, procedures, index)

    def _generate_basic_program_doc(self, program_name: str, prog_summary: Dict,
                                    prog_details: Dict, dependencies: Dict,
                                    data_flows: Dict, procedures: List[Dict], index: int) -> str:
        """Generate basic documentation without LLM (fallback)"""

        md = f"""### 3.{index} {program_name}

#### Program Information

| Property | Value |
|----------|-------|
| **Program Name** | {program_name} |
| **Domain** | {prog_summary.get('domain', 'N/A')} |
| **Complexity Score** | {prog_summary.get('complexity') or 0} |
| **Lines of Code** | {prog_summary.get('loc') or 0} |
| **Description** | {prog_details.get('description', 'No description available')} |

#### Purpose and Overview

{prog_details.get('description', 'This program is part of the COBOL system.')}

"""

        # Dependencies
        calls = dependencies.get('calls', [])
        called_by = dependencies.get('called_by', [])

        if calls or called_by:
            md += "#### Dependencies\n\n"

            if calls:
                md += f"**Calls {len(calls)} programs:**\n\n"
                for dep in calls[:10]:
                    md += f"- `{dep['program']}` ({dep.get('domain', 'N/A')})\n"
                if len(calls) > 10:
                    md += f"\n*...and {len(calls) - 10} more*\n"
                md += "\n"

            if called_by:
                md += f"**Called by {len(called_by)} programs:**\n\n"
                for dep in called_by[:10]:
                    md += f"- `{dep['program']}` ({dep.get('domain', 'N/A')})\n"
                if len(called_by) > 10:
                    md += f"\n*...and {len(called_by) - 10} more*\n"
                md += "\n"

        # Data Operations
        if data_flows:
            md += "#### Data Operations\n\n"
            md += "| File Name | Operations | Description |\n"
            md += "|-----------|------------|-------------|\n"

            for file_name, file_info in data_flows.items():
                ops = ", ".join(set(file_info['operations']))
                desc = file_info.get('description', 'N/A')
                md += f"| {file_name} | {ops} | {desc} |\n"
            md += "\n"

        # Procedures
        if procedures:
            md += f"#### Internal Structure\n\n"
            md += f"This program contains {len(procedures)} procedures/paragraphs:\n\n"

            for proc in procedures[:20]:
                md += f"- **{proc['name']}**"
                if proc.get('description'):
                    md += f": {proc['description']}"
                md += "\n"

            if len(procedures) > 20:
                md += f"\n*...and {len(procedures) - 20} more procedures*\n"
            md += "\n"

        # Technical Notes
        complexity = prog_summary.get('complexity') or 0
        if complexity > 70:
            md += "#### Technical Notes\n\n"
            md += f"⚠️ **High Complexity Warning**: This program has a complexity score of {complexity}, "
            md += "which may make it difficult to maintain. Consider refactoring or adding comprehensive tests.\n\n"

        return md

    def _generate_program_documentation(self, data: Dict) -> str:
        """Generate documentation for a single program (detailed mode) with full skeleton"""

        prog_info = data.get('program_info', {})
        program_name = prog_info.get('name', 'Unknown')
        copybooks = data.get('copybooks', [])
        jobs = data.get('jobs', [])
        data_flows = data.get('data_flows', {})
        dependencies = data.get('dependencies', {})
        source_code = prog_info.get('source_code') or ''

        # Determine enrichment status
        is_enriched = prog_info.get('domain') not in ['Not Enriched', None]
        complexity = prog_info.get('complexity_score') or 0

        # Count source-referenced files via SELECT...ASSIGN
        source_file_count = len(re.findall(r'SELECT\s+[\w-]+\s+ASSIGN', source_code.upper())) if source_code else 0

        prog_summary = {
            'program_name': program_name,
            'domain': prog_info.get('domain'),
            'complexity': complexity,
            'loc': prog_info.get('loc'),
            'source_code': source_code
        }

        # --- Summary table ---
        enrichment_label = "Yes" if is_enriched else "No"
        enrichment_warning = ""
        if not is_enriched:
            enrichment_warning = f"""
> **Enrichment Note:** The knowledge graph does not contain enriched metadata (business_logic, description) for this program. The documentation below was generated by analyzing the COBOL source code directly.
"""

        complexity_label = f"{complexity} (HIGH)" if complexity >= 70 else f"{complexity} ({'MEDIUM' if complexity >= 30 else 'LOW'})"

        md = f"""# COBOL Program Technical Documentation — {program_name}

**Document Type:** Program Detail Specification
**Generated:** {data['generation_date']}
**Program:** {program_name}
**Domain:** {prog_info.get('domain') or 'Not Enriched'}
**Enrichment Status:** {'Enriched' if is_enriched else 'Not Enriched — documentation inferred from source code + knowledge graph'}

---

## Program Summary

| Property | Value |
|----------|-------|
| Program Name | {program_name} |
| Domain | {prog_info.get('domain') or 'Not Enriched'} |
| Complexity Score | {complexity_label} |
| Lines of Code | {prog_info.get('loc') or 0} |
| Source Code Available | {'Yes' if source_code else 'No'} |
| Enriched | {enrichment_label} |
| Outbound Calls | {len(dependencies.get('calls', []))} |
| Inbound Calls | {len(dependencies.get('called_by', []))} |
| Data Files (graph) | {len(data_flows)} |
| Data Files (source) | {source_file_count} |
| Copybooks | {len(copybooks)} |
| JCL Jobs | {len(jobs)} |

{enrichment_warning}
---

"""

        # --- LLM-generated body ---
        md += self._generate_single_program_doc(prog_summary, 1, copybooks=copybooks, jobs=jobs)

        # --- Footer ---
        md += f"""

---

**END OF DOCUMENT**

*Generated by COBOL Agentic Knowledge Graph*
*Source: Neo4j Knowledge Graph + COBOL source code analysis*
*Generated: {data['generation_date']}*
"""

        return md

    def _export_file(self, doc_type: str, doc_format: str, content: str, program_name: str = None) -> Path:
        """Export document to file"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if program_name:
            filename = f"TechDoc_{program_name}_{timestamp}.{self._get_file_extension(doc_format)}"
        else:
            filename = f"TechDoc_{doc_type}_{timestamp}.{self._get_file_extension(doc_format)}"

        file_path = self.exports_dir / filename
        file_path.write_text(content, encoding='utf-8')

        logger.info(f"💾 Exported technical documentation to: {file_path}")
        return file_path

    def _get_file_extension(self, doc_format: str) -> str:
        """Get file extension for format"""
        extensions = {
            'markdown': 'md',
            'docx': 'docx',
            'word': 'docx',
            'pdf': 'pdf'
        }
        return extensions.get(doc_format, 'txt')

    def _export_as_docx(self, doc_type: str, markdown_content: str, program_name: str = None) -> Path:
        """
        Convert markdown documentation to DOCX format

        Args:
            doc_type: Type of document
            markdown_content: Markdown formatted content
            program_name: Optional program name for filename

        Returns:
            Path to exported DOCX file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if program_name:
            filename = f"TechDoc_{program_name}_{timestamp}.docx"
        else:
            filename = f"TechDoc_{doc_type}_{timestamp}.docx"

        file_path = self.exports_dir / filename

        # Create DOCX document
        doc = Document()

        # Process markdown line by line
        lines = markdown_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # Main title (# Title)
            if line.startswith('# ') and not line.startswith('##'):
                title_text = line.replace('# ', '')
                title = doc.add_heading(title_text, level=0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Heading level 2 (## Heading)
            elif line.startswith('## '):
                heading_text = re.sub(r'^##\s+', '', line)
                heading_text = re.sub(r'[📋🏗️🔄🔗💼💥⚠️🎯📎]', '', heading_text).strip()
                doc.add_heading(heading_text, level=1)

            # Heading level 3 (### Heading)
            elif line.startswith('### '):
                heading_text = re.sub(r'^###\s+', '', line)
                heading_text = re.sub(r'^\d+\.\d+\s+', '', heading_text)  # Remove numbering like "3.1"
                doc.add_heading(heading_text, level=2)

            # Heading level 4 (#### Heading)
            elif line.startswith('#### '):
                heading_text = re.sub(r'^####\s+', '', line)
                doc.add_heading(heading_text, level=3)

            # Table detection
            elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = [line]
                i += 1
                # Collect all table lines
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                i -= 1

                # Parse and create table
                self._add_table_to_doc(doc, table_lines)

            # Blockquote / Warning (> text)
            elif line.startswith('>'):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_lines.append(lines[i].strip().replace('> ', '').replace('>', ''))
                    i += 1
                i -= 1

                para = doc.add_paragraph('\n'.join(quote_lines))
                para.paragraph_format.left_indent = Inches(0.5)
                para.runs[0].font.italic = True
                para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Bullet list (- item or * item)
            elif line.startswith('- ') or line.startswith('* '):
                list_text = re.sub(r'^[-\*]\s+', '', line)
                list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)  # Remove bold markers
                list_text = re.sub(r'`(.*?)`', r'\1', list_text)  # Remove code markers
                doc.add_paragraph(list_text, style='List Bullet')

            # Horizontal rule (---)
            elif line.startswith('---'):
                doc.add_paragraph()  # Just add spacing

            # Code block (```)
            elif line.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    para = doc.add_paragraph('\n'.join(code_lines))
                    para.runs[0].font.name = 'Courier New'
                    para.runs[0].font.size = Pt(9)
                    para.paragraph_format.left_indent = Inches(0.5)

            # Bold text (**text** or __text__)
            elif '**' in line or '__' in line:
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)

            # Regular paragraph
            else:
                # Clean up markdown formatting
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Remove bold
                text = re.sub(r'`(.*?)`', r'\1', text)  # Remove code
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links
                if text.strip():
                    doc.add_paragraph(text)

            i += 1

        # Save document
        doc.save(str(file_path))
        logger.info(f"💾 Exported DOCX documentation to: {file_path}")

        return file_path

    def _add_table_to_doc(self, doc, table_lines: List[str]):
        """Add a markdown table to the DOCX document"""
        if len(table_lines) < 2:
            return

        # Parse table
        rows = []
        for line in table_lines:
            if '---' in line and all(c in '|-: ' for c in line):
                continue  # Skip separator line
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'

        # Fill table
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                # Clean markdown formatting
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_data)
                text = re.sub(r'`(.*?)`', r'\1', text)
                text = re.sub(r'[🔴🟡🟢⚠️✅❌📊💰💵💳]', '', text).strip()
                cell.text = text

                # Bold header row
                if i == 0:
                    cell.paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()  # Add spacing after table

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with markdown formatting to a paragraph"""
        # Simple bold detection
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part.strip('*'))
                run.font.bold = True
            else:
                paragraph.add_run(part)


# Create singleton instance
document_generator_agent = DocumentGeneratorAgent()


# Wrapper function for LangGraph (if needed)
def document_generator_agent_node(state: Dict[str, Any]) -> Dict[str, Any]:
    """LangGraph node wrapper"""
    return document_generator_agent.process(state)
